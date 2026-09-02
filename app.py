import streamlit as st
from docx import Document
import io
import pypdf
import google.generativeai as genai

# Configuration de la page
st.set_page_config(
    page_title="Analyseur DCE & Revue d'Offre - Eiffage",
    page_icon="🏗️",
    layout="wide",
)

st.title("🏗️ Assistant IA (Gemini 3.5) - Analyse de DCE & Revue d'Offre")
st.markdown("Dépose tes pièces de DCE. Gemini 3.5 analyse les documents point par point et te génère un tableau Word structuré.")

st.divider()

# Saisie de la clé API Gemini dans la barre latérale
api_key = st.sidebar.text_input("Clé API Gemini", type="password", help="Entre ta clé API Google AI Studio.")

# Zone de dépôt des pièces du DCE
dce_files = st.file_uploader(
    "📁 Dépose tes pièces du DCE (PDF, CCTP, RC, etc.)", 
    type=["pdf", "docx", "txt"], 
    accept_multiple_files=True,
    help="Glisse l'ensemble des pièces de l'appel d'offres ici."
)

st.divider()

if st.button("🚀 Lancer l'analyse du DCE et générer le tableau Word", type="primary", use_container_width=True):
    if not dce_files:
        st.warning("⚠️ Veuillez déposer au moins une pièce du DCE.")
    elif not api_key:
        st.error("⚠️ Veuillez renseigner votre clé API Gemini dans la barre latérale.")
    else:
        # Création de la barre de progression visuelle
        progress_bar = st.progress(0)
        status_text = st.empty()

        try:
            # ÉTAPE 1 : Extraction du texte (30%)
            status_text.text("🔄 Étape 1/3 : Extraction du texte des pièces du DCE...")
            progress_bar.progress(10)
            
            texte_dce_total = ""
            total_files = len(dce_files)
            
            for index, uploaded_file in enumerate(dce_files):
                if uploaded_file.name.endswith(".pdf"):
                    reader = pypdf.PdfReader(uploaded_file)
                    for i, page in enumerate(reader.pages):
                        texte_dce_total += f"\n--- Document: {uploaded_file.name} (Page {i+1}) ---\n"
                        texte_dce_total += page.extract_text() or ""
                else:
                    texte_dce_total += f"\n--- Document: {uploaded_file.name} ---\n"
                    texte_dce_total += uploaded_file.read().decode("utf-8", errors="ignore")
                
                current_progress = int(10 + ((index + 1) / total_files) * 20)
                progress_bar.progress(current_progress)

            # ÉTAPE 2 : Analyse par Gemini 3.5 (70%)
            status_text.text("🤖 Étape 2/3 : Analyse approfondie point par point par Gemini 3.5...")
            progress_bar.progress(40)

            genai.configure(api_key=api_key)
            
            prompt_systeme = (
                "Tu es un expert en réponse aux appels d'offres de maintenance et performance énergétique pour Eiffage Énergie Systèmes. "
                "Analyse le DCE fourni et réponds aux 67 points de la grille de revue d'offre. "
                "Pour chaque point, commence ta ligne exactement par le numéro au format 'POINT X :' (ex: 'POINT 1 :', 'POINT 2 :', etc.) "
                "suivi de l'analyse détaillée, des informations trouvées, et de la page de référence du document."
            )

            # Utilisation explicite du modèle gemini-3.5-flash
            model = genai.GenerativeModel('gemini-3.5-flash')
            
            progress_bar.progress(60)
            
            response = model.generate_content(
                [
                    prompt_systeme,
                    f"Voici le contenu complet du DCE :\n{texte_dce_total}"
                ],
                request_options={"timeout": 120}  # Tolérance de 120 secondes
            )
            
            analyse_resultat = response.text
            progress_bar.progress(85)

            # ÉTAPE 3 : Création du document Word (100%)
            status_text.text("📝 Étape 3/3 : Mise en forme et création du tableau Word...")
            
            doc = Document()
            doc.add_heading("FICHE DE REVUE D’OFFRE - EXPLOITATION MAINTENANCE", level=1)
            doc.add_paragraph("Tableau d'analyse automatisée des pièces du DCE - Eiffage Énergie Systèmes.")
            
            doc.add_heading("Grille d'analyse détaillée", level=2)

            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Numéro du Point"
            hdr_cells[1].text = "Analyse et Réponse extraite du DCE"

            lignes = analyse_resultat.split("\n")
            for ligne in lignes:
                if "POINT" in ligne.upper() and ":" in ligne:
                    parties = ligne.split(":", 1)
                    row_cells = table.add_row().cells
                    row_cells[0].text = parties[0].strip()
                    row_cells[1].text = parties[1].strip() if len(parties) > 1 else ""

            if len(table.rows) == 1:
                row_cells = table.add_row().cells
                row_cells[0].text = "Synthèse globale"
                row_cells[1].text = analyse_resultat

            output_io = io.BytesIO()
            doc.save(output_io)
            output_io.seek(0)

            progress_bar.progress(100)
            status_text.text("✅ Analyse terminée avec succès !")

            st.success("✅ Fiche de revue d'offre structurée en tableau générée avec succès !")
            
            st.download_button(
                label="📥 Télécharger la Fiche Word structurée (.docx)",
                data=output_io,
                file_name="Fiche_Revue_Offre_Gemini_3.5.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )

        except Exception as e:
            st.error(f"Une erreur est survenue lors du traitement : {e}")
